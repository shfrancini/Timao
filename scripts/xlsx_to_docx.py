import sys
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def compute_numbering_path(previous_path, current_labels):
    path = []
    for i, label in enumerate(current_labels):
        if not label:
            break
        if i >= len(previous_path):
            path.append(1)
        elif label == previous_path[i][1]:
            path.append(previous_path[i][0])
        else:
            path.append(previous_path[i][0] + 1)
    return [(n, label) for n, label in zip(path, current_labels)]


def format_numbering_path(numbering_path):
    return '.'.join(str(n) for n, _ in numbering_path if _) + '.'


def main(input_xlsx, output_docx):
    wb = openpyxl.load_workbook(input_xlsx)
    sheet = wb.active

    headers = [cell.value for cell in sheet[1]]
    rows = list(sheet.iter_rows(min_row=2, values_only=True))

    hierarchy_cols = [h for h in headers if h and h.startswith("H")]
    gpt_cols = [h for h in headers if h and "gpt" in h.lower()]
    cctp_col = headers.index("CCTP") if "CCTP" in headers else None

    if cctp_col is None:
        print("❌ 'CCTP' column not found.")
        sys.exit(1)

    doc = Document()
    previous_path = []

    for row in rows:
        current_labels = [row[headers.index(h)] for h in hierarchy_cols]
        numbering_path = compute_numbering_path(previous_path, current_labels)
        previous_path = numbering_path

        for i, (num, label) in enumerate(numbering_path):
            if not label:
                continue
            level = i + 1
            numbered_title = f"{format_numbering_path(numbering_path[:level])} {label}"
            doc.add_heading(numbered_title, level=level)

        # Add CCTP content only if it exists
        cctp_value = row[cctp_col]
        if cctp_value and str(cctp_value).strip():
            doc.add_paragraph("**CCTP**", style="Intense Quote")
            para = doc.add_paragraph(str(cctp_value).strip())
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # Add GPT responses
        for gpt_col in gpt_cols:
            value = row[headers.index(gpt_col)]
            if value and str(value).strip():
                doc.add_paragraph(f"**{gpt_col.replace('_', ' ').capitalize()}**", style="Intense Quote")
                para = doc.add_paragraph(str(value))
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = para.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(11)

        doc.add_paragraph("")

    doc.save(output_docx)
    print(f"✅ DOCX saved to {output_docx}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("❌ Usage: python3 xlsx_to_docx.py input.xlsx output.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    main(input_file, output_file)
